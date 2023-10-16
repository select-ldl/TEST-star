# -*- coding:utf-8 -*-
# @Author: start.liu
# @Email: start.liu@stert.liu
# @Date: 2022/6/22 20:15
# @File: pase
# @Project: gutuibai



import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx
import requests
from bs4 import BeautifulSoup
class Myswel():
    def __init__(self):
        self.filenames()
        self.reade()
        self.chang()
        self.write()
    def filenames(self):
        "获取当前文件"
        for path,list,filelist in os.walk(os.getcwd()):
            # print(path,list,filelist)
            for pys in filelist:
                "获取对应文件的文件名字"
                # print(os.path.splitext(pys))
                if os.path.splitext(pys)[1]=='.docx':
                    "获取名字"
                    self.name=os.path.splitext(pys)[0]
                    #入经
                    self.files=os.path.abspath('%s.docx'%self.name)
    def reade(self):
        "读取文件"
        self.wenzhang = []
        doc = docx.Document(self.files)
        # print(files)
        for para in doc.paragraphs:
            filerede=para.text

            self.wenzhang.append(filerede)
        # print(self.wenzhang)

    # print(re.split('.',str(pys)))

    def write(self):
        "写入文件"
        # 在WD_PARAGRAPH_ALIGNMENT可以实现LEFT、RIGHT、CENTER、JUSTY和DISTRIBUTE等5种对齐方式
        # WD_PARAGRAPH_ALIGNMENT.LEFT # 左对齐
        # WD_PARAGRAPH_ALIGNMENT.CENTER # 居中对齐
        # WD_PARAGRAPH_ALIGNMENT.RIGHT # 右对齐
        # WD_PARAGRAPH_ALIGNMENT.JUSTIFY # 两端对齐
        # WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE # 分散对齐

        file=docx.Document()
        # =WD_PARAGRAPH_ALIGNMENT.CENTER

        file.add_heading(text="%s"%self.telit,level=1)
        for i in self.zetd:
            text = ''
            for j in i:
                j+="."
                text+=j
            file.add_paragraph(text=text+'\n', style=None)
        file.save('%s/%sEN.docx' % (os.path.abspath('/'), self.name))
        print("写完了？")
    def huoqu(self,a):
        urls = 'https://translate.google.cn/m?sl=auto&tl=en&hl=zh-CN'
        dict = {'q': '%s' % a}
        # print(r.text)
        header = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.89 Safari/537.36 SLBrowser/7.0.0.12151 SLBChan/30'}
        reqe = requests.get(url=urls, params=dict)
        reqe.encoding = 'utf-8'
        # print(reqe.url)
        sopu = BeautifulSoup(reqe.text, 'lxml')
        mu = sopu.select('body > div.root-container > div.result-container')
        for nsk in mu:
            self.netext=nsk.text
        return self.netext

    def chang(self):
        '处理文件'
        retd=[]
        self.zetd=[]
        for i in self.wenzhang:
            retd=[]
            for j in i.split('。'):
                tex=self.huoqu(j)
                if tex != '':
                    retd.append(tex)
            if retd!=['']:
                self.zetd.append(retd)
        tt=self.zetd[0]
        for ise in tt:
            self.telit=ise
        del self.zetd[0]
        print(self.telit)
        print("翻译完了？")
if __name__ == '__main__':
    Myswel()
    # a= ['Ten Minutes Introductory Financial Statements']
    # print(a)

